"""Staged document extraction that explains itself.

The previous ``import_document`` returned a bare string and raised one generic
error for every failure mode, so "no text could be extracted" covered scanned
PDFs, corrupt files, and empty files alike. This module keeps the extraction
attempt and its *diagnosis* together, so the UI can say what actually went
wrong and what the user can do about it.

Stages:

1. ``detect_kind``  — magic bytes first, extension only as a hint.
2. ``extract``      — per-format, with fallbacks; never raises for empty output.
3. ``assess``       — classify the result (OK / IMAGE_ONLY / THIN / ...).
4. ``normalise``    — repair the artefacts PDF text extraction always produces.

``load_document`` runs all four. Callers that want the raw text without
diagnosis can still use ``documents.importers.import_document``.
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from enum import StrEnum
from pathlib import Path

from aptiordesk.core.errors import DocumentError

log = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

# Below this, a "successful" extraction is almost certainly a failure wearing a
# success mask — a cover page, a header, or one stray text object on a scan.
THIN_TEXT_THRESHOLD = 200


class Diagnosis(StrEnum):
    OK = "ok"
    THIN = "thin"
    IMAGE_ONLY = "image_only"
    EMPTY = "empty"
    ENCRYPTED = "encrypted"
    CORRUPT = "corrupt"
    UNSUPPORTED = "unsupported"


#: Whether the extraction is usable enough to send to the AI at all.
_USABLE = {Diagnosis.OK, Diagnosis.THIN}

_ADVICE = {
    Diagnosis.IMAGE_ONLY: (
        "This PDF has no text layer — the pages are images, which usually means "
        "it was scanned or exported as pictures. AptiorDesk cannot read text from "
        "images. Export a text-based PDF from the original document (Word, Google "
        "Docs, LaTeX), or run it through an OCR tool first."
    ),
    Diagnosis.EMPTY: "This document contains no text at all.",
    Diagnosis.ENCRYPTED: (
        "This PDF is password-protected. Open it in a PDF reader, save an "
        "unprotected copy, and import that instead."
    ),
    Diagnosis.CORRUPT: (
        "This file could not be parsed. It may be damaged, or it may not really "
        "be the format its extension claims."
    ),
    Diagnosis.THIN: (
        "Only a small amount of text was found. If your resume is longer than "
        "this, the file may be partly image-based and the extraction is "
        "incomplete — check the extracted text before continuing."
    ),
}


@dataclass
class ExtractionResult:
    """Everything known about one extraction attempt, success or failure."""

    text: str = ""
    raw_text: str = ""
    diagnosis: Diagnosis = Diagnosis.OK
    kind: str = ""
    filename: str = ""
    page_count: int = 0
    pages_with_text: int = 0
    method: str = ""
    warnings: list[str] = field(default_factory=list)
    detail: str = ""

    @property
    def ok(self) -> bool:
        """True when there is enough text to attempt structuring."""
        return self.diagnosis in _USABLE and bool(self.text.strip())

    @property
    def char_count(self) -> int:
        return len(self.text)

    def message(self) -> str:
        """A user-facing explanation. Empty when the extraction is clean."""
        if self.diagnosis is Diagnosis.OK:
            return ""
        advice = _ADVICE.get(self.diagnosis, "This document could not be read.")
        if self.diagnosis is Diagnosis.IMAGE_ONLY and self.page_count:
            return f"No text layer found in any of the {self.page_count} page(s). " + advice
        return advice

    def raise_if_unusable(self) -> None:
        if not self.ok:
            raise DocumentError(
                self.message() or "This document could not be read.", detail=self.detail
            )


# --- stage 1: identify --------------------------------------------------------

_MAGIC: list[tuple[bytes, str]] = [
    (b"%PDF", "pdf"),
    (b"PK\x03\x04", "zip"),  # DOCX and every other OOXML/zip file
    (b"\xd0\xcf\x11\xe0", "doc"),  # legacy OLE2 .doc — detected so we can say so
    (b"{\\rtf", "rtf"),
]


def detect_kind(path: Path) -> str:
    """Identify the format from content, falling back to the extension.

    Content wins: a .docx that is really a PDF is a PDF, and a .pdf that is
    really an OLE2 .doc gets a precise message instead of "corrupt".
    """
    with path.open("rb") as handle:
        head = handle.read(8)
    for magic, kind in _MAGIC:
        if head.startswith(magic):
            if kind == "zip":
                return "docx" if path.suffix.lower() == ".docx" else "zip"
            return kind
    extension = path.suffix.lower()
    if extension in {".txt", ".md", ".markdown"}:
        return "text"
    return "unknown"


def validate(path: Path) -> None:
    """Cheap checks that do not require opening the document."""
    if not path.is_file():
        raise DocumentError(f"File not found: {path.name}")
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentError(
            f"Unsupported file type '{extension or path.name}'. "
            "AptiorDesk reads PDF, DOCX, TXT, and Markdown."
        )
    size = path.stat().st_size
    if size == 0:
        raise DocumentError("The file is empty (0 bytes).")
    if size > MAX_FILE_SIZE:
        raise DocumentError(f"File is too large ({size / 1_048_576:.1f} MB). The limit is 10 MB.")


# --- stage 2: extract ---------------------------------------------------------


def extract(path: Path, kind: str) -> ExtractionResult:
    """Pull text out of the document. Returns a diagnosis rather than raising
    for content problems; only genuinely unreadable files raise."""
    result = ExtractionResult(kind=kind, filename=path.name)
    if kind == "pdf":
        _extract_pdf(path, result)
    elif kind == "docx":
        _extract_docx(path, result)
    elif kind == "text":
        _extract_text(path, result)
    elif kind == "doc":
        result.diagnosis = Diagnosis.UNSUPPORTED
        result.detail = "legacy OLE2 .doc container"
        raise DocumentError(
            "This is a legacy Word .doc file, not a .docx. Open it in Word and "
            "save as .docx or PDF, then import that."
        )
    else:
        result.diagnosis = Diagnosis.CORRUPT
        result.detail = f"unrecognised container: {kind}"
    return result


def _extract_pdf(path: Path, result: ExtractionResult) -> None:
    pages: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            # An empty user password is common and harmless; try it before giving up.
            try:
                opened = reader.decrypt("")
            except Exception:
                opened = 0
            if not opened:
                result.diagnosis = Diagnosis.ENCRYPTED
                return
        result.page_count = len(reader.pages)
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:  # one bad page must not lose the rest
                pages.append("")
                result.warnings.append(f"A page could not be read ({exc.__class__.__name__}).")
        result.method = "pypdf"
    except Exception as exc:
        log.warning("pypdf failed on %s: %s — falling back to pdfplumber", path.name, exc)
        result.detail = str(exc)

    # pdfplumber does better on multi-column and table-heavy layouts, so try it
    # whenever pypdf produced nothing or produced suspiciously little.
    if sum(len(p.strip()) for p in pages) < THIN_TEXT_THRESHOLD:
        plumber_pages = _pdfplumber_pages(path, result)
        if plumber_pages is not None and sum(len(p.strip()) for p in plumber_pages) > sum(
            len(p.strip()) for p in pages
        ):
            pages = plumber_pages
            result.method = "pdfplumber"
            result.page_count = len(pages)

    if not pages and not result.page_count:
        result.diagnosis = Diagnosis.CORRUPT
        return

    result.pages_with_text = sum(1 for p in pages if p.strip())
    result.raw_text = "\n\n".join(pages).strip()
    if result.page_count and result.pages_with_text == 0:
        # Pages exist, none carry text: the defining signature of a scan.
        result.diagnosis = Diagnosis.IMAGE_ONLY
    elif result.page_count and result.pages_with_text < result.page_count:
        result.warnings.append(
            f"{result.page_count - result.pages_with_text} of {result.page_count} "
            "page(s) had no text layer and were skipped."
        )


def _pdfplumber_pages(path: Path, result: ExtractionResult) -> list[str] | None:
    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return [page.extract_text() or "" for page in pdf.pages]
    except Exception as exc:
        log.warning("pdfplumber failed on %s: %s", path.name, exc)
        result.detail = result.detail or str(exc)
        return None


def _extract_docx(path: Path, result: ExtractionResult) -> None:
    try:
        import docx

        document = docx.Document(str(path))
    except Exception as exc:
        result.diagnosis = Diagnosis.CORRUPT
        result.detail = str(exc)
        return

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Many resumes lay out dates/roles in two-column tables; a tab keeps
            # the columns distinguishable without inventing a bullet structure.
            if any(cells):
                parts.append("\t".join(cells))
    # Headers and footers often hold the contact block in designed templates.
    for section in document.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
    result.method = "python-docx"
    result.page_count = 1
    result.raw_text = "\n".join(parts).strip()
    result.pages_with_text = 1 if result.raw_text else 0


def _extract_text(path: Path, result: ExtractionResult) -> None:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            result.raw_text = data.decode(encoding).strip()
            result.method = f"text/{encoding}"
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        result.diagnosis = Diagnosis.CORRUPT
        result.detail = "could not decode with any known encoding"
        return
    result.page_count = 1
    result.pages_with_text = 1 if result.raw_text else 0


# --- stage 3: assess ----------------------------------------------------------


def assess(result: ExtractionResult) -> ExtractionResult:
    """Classify an extraction whose diagnosis is not already decided."""
    if result.diagnosis is not Diagnosis.OK:
        return result
    stripped = result.text.strip() or result.raw_text.strip()
    if not stripped:
        result.diagnosis = Diagnosis.IMAGE_ONLY if result.page_count else Diagnosis.EMPTY
    elif len(stripped) < THIN_TEXT_THRESHOLD:
        result.diagnosis = Diagnosis.THIN
    return result


# --- stage 4: normalise -------------------------------------------------------

_LIGATURES = {
    "ﬀ": "ff",
    "ﬁ": "fi",
    "ﬂ": "fl",
    "ﬃ": "ffi",
    "ﬄ": "ffl",
}
# Bullet glyphs PDF producers emit; all mean "list item".
_BULLETS = "•·▪◦‣⁃∙◾▸►❖✦*"
_HYPHEN_BREAK = re.compile(r"(\w)[-‐‑]\n\s*(\w)")
_MANY_BLANKS = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")
_MANY_SPACES = re.compile(r"[ \t]{3,}")
_PAGE_MARKER = re.compile(
    r"^\s*(?:page\s+)?\d+\s*(?:/|of)\s*\d+\s*$|^\s*-\s*\d+\s*-\s*$", re.IGNORECASE
)


def normalise(text: str) -> str:
    """Repair the artefacts PDF and DOCX extraction reliably introduce.

    Deliberately conservative: it removes noise that would confuse the model
    but never rewrites the candidate's own wording.
    """
    if not text:
        return ""
    text = unicodedata.normalize("NFC", text)
    for ligature, replacement in _LIGATURES.items():
        text = text.replace(ligature, replacement)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace("​", "")
    # Curly quotes and dashes carry no meaning here and cost tokens.
    text = (
        text.replace("‘", "'")
        .replace("’", "'")
        .replace("“", '"')
        .replace("”", '"')
        .replace("–", "-")
        .replace("—", "-")
    )
    # "engi-\nneering" -> "engineering". Only across a line break, so real
    # hyphenated words ("full-stack") are untouched.
    text = _HYPHEN_BREAK.sub(r"\1\2", text)

    lines: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if _PAGE_MARKER.match(stripped):
            continue  # "Page 2 of 3" footers
        for bullet in _BULLETS:
            if stripped.startswith(bullet):
                stripped = "- " + stripped[len(bullet) :].lstrip()
                break
        lines.append(_MANY_SPACES.sub("  ", stripped))

    text = "\n".join(lines)
    text = _drop_repeated_lines(text)
    text = _TRAILING_WS.sub("\n", text)
    text = _MANY_BLANKS.sub("\n\n", text)
    return text.strip()


def _drop_repeated_lines(text: str) -> str:
    """Remove running headers/footers — short lines repeated on most pages.

    Requires 3+ occurrences so a resume that legitimately repeats a short line
    twice keeps both.
    """
    lines = text.split("\n")
    counts: dict[str, int] = {}
    for line in lines:
        stripped = line.strip()
        if 0 < len(stripped) <= 60:
            counts[stripped] = counts.get(stripped, 0) + 1
    repeated = {line for line, n in counts.items() if n >= 3}
    if not repeated:
        return text
    return "\n".join(line for line in lines if line.strip() not in repeated)


# --- the whole pipeline -------------------------------------------------------


def load_document(path: str | Path) -> ExtractionResult:
    """Validate, extract, normalise, and diagnose a document in one call.

    Raises ``DocumentError`` only for files that cannot be opened at all;
    content problems (scans, near-empty files) come back as a diagnosis so the
    caller can show the raw text alongside the explanation.
    """
    path = Path(path)
    validate(path)
    kind = detect_kind(path)
    if kind in {"zip", "unknown", "rtf"}:
        raise DocumentError(
            f"'{path.name}' is not a readable {path.suffix.lstrip('.').upper()} file. "
            "Its contents do not match its extension."
        )
    result = extract(path, kind)
    result.text = normalise(result.raw_text)
    result = assess(result)
    log.info(
        "Imported %s [%s via %s] -> %d chars, diagnosis=%s",
        path.name,
        kind,
        result.method or "?",
        result.char_count,
        result.diagnosis,
    )
    return result
