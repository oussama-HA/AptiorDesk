"""Document import: extract plain text from PDF / DOCX / TXT / MD files.

All uploads are untrusted input: size-capped, extension- and magic-byte-
validated, and parsing failures surface as ``DocumentError`` with a message
safe to show the user. Extracted text is *raw material* — the user always
reviews the structured extraction before anything is saved.
"""

from __future__ import annotations

import logging
from pathlib import Path

from aptiordesk.core.errors import DocumentError

log = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

_MAGIC = {
    ".pdf": b"%PDF",
    ".docx": b"PK\x03\x04",  # OOXML zip container
}


def import_document(path: str | Path) -> str:
    """Return extracted plain text for a supported document."""
    path = Path(path)
    _validate(path)
    extension = path.suffix.lower()
    if extension == ".pdf":
        text = _extract_pdf(path)
    elif extension == ".docx":
        text = _extract_docx(path)
    else:
        text = _extract_text(path)

    text = text.strip()
    if not text:
        raise DocumentError(
            "No text could be extracted from this file. If it is a scanned PDF, "
            "it contains images rather than text — export a text-based version instead."
        )
    log.info("Imported %s (%d chars)", path.name, len(text))
    return text


def _validate(path: Path) -> None:
    if not path.is_file():
        raise DocumentError(f"File not found: {path.name}")
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentError(
            f"Unsupported file type '{extension}'. Supported: PDF, DOCX, TXT, Markdown."
        )
    size = path.stat().st_size
    if size == 0:
        raise DocumentError("The file is empty.")
    if size > MAX_FILE_SIZE:
        raise DocumentError(f"File is too large ({size / 1_048_576:.1f} MB). The limit is 10 MB.")
    magic = _MAGIC.get(extension)
    if magic:
        with path.open("rb") as handle:
            if not handle.read(len(magic)).startswith(magic):
                raise DocumentError(
                    f"This file does not look like a valid {extension[1:].upper()} file."
                )


def _extract_pdf(path: Path) -> str:
    # pypdf first (fast, pure); pdfplumber as fallback for layout-heavy PDFs.
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        if reader.is_encrypted:
            raise DocumentError("This PDF is password-protected. Remove the password first.")
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        if text.strip():
            return text
    except DocumentError:
        raise
    except Exception as exc:
        log.warning("pypdf failed on %s: %s — trying pdfplumber", path.name, exc)

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            return "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as exc:
        raise DocumentError(
            "Could not read this PDF. It may be corrupted or image-only.",
            detail=str(exc),
        ) from exc


def _extract_docx(path: Path) -> str:
    try:
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        raise DocumentError(
            "Could not read this DOCX file. It may be corrupted.", detail=str(exc)
        ) from exc


def _extract_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as exc:
            raise DocumentError("Could not decode this text file.", detail=str(exc)) from exc
