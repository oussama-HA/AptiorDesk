"""Export documents to Markdown, plain text, PDF, and DOCX.

Content and formatting are separate: everything upstream produces Markdown,
and these functions render it. PDF uses Qt's own text engine (no extra
dependency, works offline); DOCX uses python-docx.

ATS-safety note: output is single-column, no tables, no text inside graphics,
no headers/footers — the formats parsers handle most reliably.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from markdown_it import MarkdownIt

from aptiordesk.core.errors import DocumentError

log = logging.getLogger(__name__)

EXPORT_FORMATS = {
    "md": "Markdown (*.md)",
    "txt": "Plain text (*.txt)",
    "pdf": "PDF (*.pdf)",
    "docx": "Word document (*.docx)",
}

_PDF_CSS = """
body { font-family: Georgia, 'Times New Roman', serif; font-size: 11pt;
       color: #000; line-height: 1.45; }
h1 { font-size: 18pt; margin: 0 0 2pt 0; }
h2 { font-size: 13pt; margin: 14pt 0 4pt 0; border-bottom: 1px solid #999; }
h3 { font-size: 11.5pt; margin: 10pt 0 2pt 0; }
p, li { font-size: 11pt; }
ul { margin: 2pt 0 2pt 14pt; }
em { color: #333; }
"""


def export_markdown(markdown: str, path: str | Path) -> Path:
    path = Path(path)
    path.write_text(markdown, encoding="utf-8")
    return path


def export_text(markdown: str, path: str | Path) -> Path:
    path = Path(path)
    path.write_text(markdown_to_plain(markdown), encoding="utf-8")
    return path


def export_pdf(markdown: str, path: str | Path) -> Path:
    """Render Markdown to PDF with Qt. Requires a QApplication (the app has
    one; tests create one via qapp)."""
    from PySide6.QtCore import QMarginsF
    from PySide6.QtGui import QPageLayout, QPageSize, QPdfWriter, QTextDocument

    path = Path(path)
    html = MarkdownIt().render(markdown)
    document = QTextDocument()
    document.setDefaultStyleSheet(_PDF_CSS)
    document.setHtml(f"<html><body>{html}</body></html>")

    writer = QPdfWriter(str(path))
    writer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
    writer.setPageMargins(QMarginsF(18, 16, 18, 16), QPageLayout.Unit.Millimeter)
    writer.setResolution(300)
    document.setPageSize(writer.pageLayout().paintRectPixels(writer.resolution()).size().toSizeF())
    document.print_(writer)
    if not path.exists() or path.stat().st_size == 0:
        raise DocumentError("PDF export produced an empty file.")
    return path


def export_docx(markdown: str, path: str | Path) -> Path:
    """Render Markdown to DOCX. Supports headings, bullets, bold/italic runs."""
    import docx

    path = Path(path)
    document = docx.Document()
    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            document.add_heading(_strip_inline(heading.group(2)), level=len(heading.group(1)))
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs(paragraph, bullet.group(1))
            continue
        paragraph = document.add_paragraph()
        _add_runs(paragraph, stripped)
    document.save(str(path))
    return path


def export_document(markdown: str, path: str | Path, fmt: str) -> Path:
    exporters = {
        "md": export_markdown,
        "txt": export_text,
        "pdf": export_pdf,
        "docx": export_docx,
    }
    if fmt not in exporters:
        raise DocumentError(f"Unsupported export format: {fmt}")
    return exporters[fmt](markdown, path)


# -- helpers ------------------------------------------------------------------

_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    """Split a line into runs so **bold** and *italic* survive into DOCX."""
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            paragraph.add_run(part[1:-1])
        else:
            paragraph.add_run(part)


def _strip_inline(text: str) -> str:
    return re.sub(r"[*`]", "", text)


def markdown_to_plain(markdown: str) -> str:
    """ATS-safe plain text: markers removed, structure preserved."""
    lines: list[str] = []
    for line in markdown.splitlines():
        stripped = line.rstrip()
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped.strip())
        if heading:
            lines.extend(["", _strip_inline(heading.group(2)).upper()])
            continue
        bullet = re.match(r"^\s*[-*]\s+(.*)$", stripped)
        if bullet:
            lines.append(f"  - {_strip_inline(bullet.group(1))}")
            continue
        lines.append(_strip_inline(stripped))
    text = "\n".join(lines)
    return re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"
