"""Cover letter generation, versioning, and export tests."""

import json
from concurrent.futures import ThreadPoolExecutor

import pytest

from aptiordesk.core.errors import DocumentError
from aptiordesk.database.models.cover_letter import CoverLetterInputs
from aptiordesk.database.models.resume import ResumeContent
from aptiordesk.documents.exporters import (
    export_document,
    export_docx,
    export_markdown,
    export_pdf,
    export_text,
    markdown_to_plain,
)
from aptiordesk.features.cover_letters.service import CoverLetterService
from aptiordesk.features.jobs.service import JobService
from aptiordesk.features.resumes.service import ResumeService
from tests.helpers import ScriptedProvider

JD = (
    "Senior Data Engineer at Initech. Requirements: 5+ years Python, Airflow, "
    "cloud data warehouses. You will design pipelines and mentor juniors. Remote."
)

LETTER_MD = """Dear Hiring Team,

When I read that Initech is rebuilding its ingestion layer, it matched the work
I have spent six years doing.

At ACME I built the Python ETL pipelines that move 2 TB a day, and I own the
Airflow deployment behind them.

I would welcome the chance to talk.

Sincerely,
Jane Roe"""

DRAFT_JSON = json.dumps(
    {
        "body_markdown": LETTER_MD,
        "selected_experiences": ["ACME ETL pipelines", "Airflow ownership"],
        "selection_rationale": "These two map directly to the posting's core asks.",
        "claims_needing_confirmation": ["Confirm the 2 TB/day figure is current."],
    }
)


def _fonts_available() -> bool:
    """Qt's offscreen platform ships no font directory; PDFs it produces have
    no embedded text layer."""
    from PySide6.QtGui import QFontDatabase

    return bool(QFontDatabase.families())


@pytest.fixture
def setup(conn):
    content = ResumeContent.model_validate(
        {
            "full_name": "Jane Roe",
            "summary": "Data engineer with 6 years of experience.",
            "experiences": [
                {
                    "title": "Data Engineer",
                    "organization": "ACME",
                    "highlights": ["Built ETL pipelines in Python processing 2 TB daily"],
                }
            ],
        }
    )
    _, version = ResumeService(conn).create_manual("Base", content)
    job = JobService(conn).create_job(JD)
    return job, version


class TestCoverLetterService:
    def test_generate_creates_version_with_rationale(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        inputs = CoverLetterInputs(tone="warm", length="short", motivation="I like their work")
        letter = service.create(job, version, inputs)
        provider = ScriptedProvider([DRAFT_JSON])

        letter_version, draft = service.generate(provider, letter, job, version, None, inputs)
        assert letter_version.version_no == 1
        assert "Initech" in letter_version.content_md
        assert draft.claims_needing_confirmation
        assert letter_version.rationale["selected_experiences"] == [
            "ACME ETL pipelines",
            "Airflow ownership",
        ]
        assert letter_version.rationale["prompt_id"] == "cover_letter"

    def test_worker_generation_never_uses_ui_sqlite_connection(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        inputs = CoverLetterInputs(tone="warm", length="short")
        letter = service.create(job, version, inputs)

        with ThreadPoolExecutor(max_workers=1) as pool:
            generated = pool.submit(
                service.generate_draft,
                ScriptedProvider([DRAFT_JSON]),
                letter,
                job,
                version,
                None,
                inputs,
            ).result()

        stored_version, draft = service.persist_generated_draft(generated)
        assert stored_version.version_no == 1
        assert draft.body_markdown == LETTER_MD
        assert service.list_versions(letter)[0].content_md == LETTER_MD

    def test_prompt_fences_untrusted_inputs(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        inputs = CoverLetterInputs()
        letter = service.create(job, version, inputs)
        provider = ScriptedProvider([DRAFT_JSON])
        service.generate(provider, letter, job, version, None, inputs)
        prompt = provider.prompts[0]
        assert "<<<BEGIN JOB DESCRIPTION>>>" in prompt
        assert "<<<BEGIN RESUME>>>" in prompt
        assert "Never invent work experience" in prompt

    def test_tone_and_length_validated(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        with pytest.raises(ValueError, match="tone"):
            service.create(job, version, CoverLetterInputs(tone="sassy"))
        with pytest.raises(ValueError, match="length"):
            service.create(job, version, CoverLetterInputs(length="epic"))

    def test_edit_creates_new_version(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        inputs = CoverLetterInputs()
        letter = service.create(job, version, inputs)
        service.generate(ScriptedProvider([DRAFT_JSON]), letter, job, version, None, inputs)
        edited = service.save_edited(letter, "My own rewritten letter.")
        assert edited.version_no == 2
        versions = service.list_versions(letter)
        assert len(versions) == 2
        assert versions[0].content_md == "My own rewritten letter."  # newest first
        assert "Initech" in versions[1].content_md  # original draft preserved
        # rationale context survives the edit, marked as inherited
        assert edited.rationale["selected_experiences"] == [
            "ACME ETL pipelines",
            "Airflow ownership",
        ]
        assert edited.rationale["inherited_from_version"] == 1

    def test_edit_without_prior_draft_has_empty_rationale(self, conn, setup):
        job, version = setup
        service = CoverLetterService(conn)
        letter = service.create(job, version, CoverLetterInputs())
        edited = service.save_edited(letter, "Hand-written letter.")
        assert edited.version_no == 1
        assert edited.rationale == {}


class TestExporters:
    def test_markdown_export(self, tmp_path):
        path = export_markdown(LETTER_MD, tmp_path / "letter.md")
        assert path.read_text(encoding="utf-8") == LETTER_MD

    def test_text_export_strips_markers(self, tmp_path):
        path = export_text("# Title\n\n- **bold** item", tmp_path / "letter.txt")
        text = path.read_text(encoding="utf-8")
        assert "TITLE" in text
        assert "**" not in text
        assert "- bold item" in text

    def test_pdf_export_is_readable(self, tmp_path, qapp):
        from pypdf import PdfReader

        path = export_pdf(
            "# Jane Roe\n\nDear Hiring Team,\n\nRegards,\nJane", tmp_path / "letter.pdf"
        )
        assert path.stat().st_size > 0
        reader = PdfReader(path)
        assert len(reader.pages) == 1
        assert path.read_bytes().startswith(b"%PDF")

        text = reader.pages[0].extract_text()
        if not _fonts_available():
            # The offscreen Qt platform ships no fonts, so glyphs cannot be
            # embedded and no text layer exists. Structure is still verified
            # above; text extraction is checked on platforms with fonts.
            pytest.skip("No fonts available on this Qt platform (offscreen)")
        assert "Jane Roe" in text
        assert "Dear Hiring Team" in text

    def test_docx_export_preserves_structure(self, tmp_path):
        import docx

        path = export_docx(
            "# Heading\n\nIntro paragraph.\n\n- First **bold** point\n- Second point",
            tmp_path / "letter.docx",
        )
        document = docx.Document(str(path))
        texts = [p.text for p in document.paragraphs]
        assert "Heading" in texts
        assert "Intro paragraph." in texts
        assert "First bold point" in texts
        bold_runs = [r.text for p in document.paragraphs for r in p.runs if r.bold]
        assert "bold" in bold_runs

    def test_export_document_dispatch(self, tmp_path, qapp):
        for fmt in ("md", "txt", "pdf", "docx"):
            path = export_document(LETTER_MD, tmp_path / f"letter.{fmt}", fmt)
            assert path.exists() and path.stat().st_size > 0

    def test_unknown_format_rejected(self, tmp_path):
        with pytest.raises(DocumentError, match="Unsupported export format"):
            export_document(LETTER_MD, tmp_path / "letter.rtf", "rtf")

    def test_markdown_to_plain_collapses_blank_runs(self):
        plain = markdown_to_plain("# A\n\n\n\nText\n\n\n- item")
        assert "\n\n\n" not in plain
